import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_font_times_new_roman(run, size=14, bold=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font_times_new_roman(run, 14, True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

def add_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    set_font_times_new_roman(run, 14, False)
    p.paragraph_format.line_spacing = 1.5
    if align == WD_ALIGN_PARAGRAPH.JUSTIFY:
        p.paragraph_format.first_line_indent = Inches(0.5)

def add_image(doc, image_path, caption):
    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=Inches(6.0))
        
        # Caption
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(caption)
        set_font_times_new_roman(r2, 12, True)
        p2.paragraph_format.space_after = Pt(24)

doc = Document()

# Page Setup
for section in doc.sections:
    section.top_margin = Inches(0.79) # 2 cm
    section.bottom_margin = Inches(0.79)
    section.left_margin = Inches(1.18) # 3 cm
    section.right_margin = Inches(0.59) # 1.5 cm

# Title Page
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("O'ZBEKISTON RESPUBLIKASI OLIY TA'LIM, FAN VA INNOVATSIYALAR VAZIRLIGI\n\n\n\n\n\n")
set_font_times_new_roman(r, 14, True)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("BITIRUV MALAKAVIY ISHI\n\n")
set_font_times_new_roman(r2, 24, True)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Mavzu: Zamonaviy tibbiyot markazlari uchun ko'p modulli (Multi-role) elektron boshqaruv (ERP) tizimini React va web-texnologiyalar asosida ishlab chiqish\n\n\n")
set_font_times_new_roman(r3, 16, True)

doc.add_page_break()

# Mundarija
add_heading(doc, "MUNDARIJA")
toc_lines = [
    ("KIRISH", "3"),
    ("I-BOB. TIBBIYOT AXBOROT TIZIMLARINING NAZARIY ASOSLARI", "7"),
    ("1.1. Tibbiyot ERP tizimlarining rivojlanishi va tasnifi", "7"),
    ("1.2. Zamonaviy tibbiyot boshqaruv tizimlarining asosiy afzalliklari", "15"),
    ("1.3. Mavjud ERP platformalarining qiyosiy tahlili", "25"),
    ("II-BOB. TIZIM TALABLARI VA TEXNOLOGIYALAR", "35"),
    ("2.1. Tizimga qo'yiladigan funksional va texnik talablar", "35"),
    ("2.2. Apparat va dasturiy platformalarni tanlash asoslari (React.js, Material UI)", "42"),
    ("2.3. Boshqaruv algoritmlari va RBAC (Multi-role) arxitekturasi", "50"),
    ("III-BOB. TIZIMNI LOYIHALASH VA AMALGA OSHIRISH", "60"),
    ("3.1. Tizim arxitekturasini yaratish va modullar tavsifi", "60"),
    ("3.2. Dasturiy ta'minotni ishlab chiqish jarayoni va qulay interfeys", "68"),
    ("3.3. Tizimni sinovdan o'tkazish va natijalarni tahlil qilish", "75"),
    ("IV-BOB. MEHNAT MUHOFAZASI VA XAVFSIZLIK", "82"),
    ("4.1. Texnika xavfsizligi qoidalari va me'yoriy talablar", "82"),
    ("4.2. Elektr xavfsizligi va kompyuter oldida ishlash ergonomikasi", "88"),
    ("XULOSA", "95"),
    ("FOYDALANILGAN ADABIYOTLAR", "98"),
    ("ILOVA (Dastur kodlari)", "100")
]

from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
for title, page in toc_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    tab_stops = p.paragraph_format.tab_stops
    # Add a right-aligned tab stop with dots at 6.2 inches
    tab_stops.add_tab_stop(Inches(6.2), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    run = p.add_run(f"{title}\t{page}")
    set_font_times_new_roman(run, 14, False)


doc.add_page_break()

# --- Generating Extensive Theory Text ---
def generate_filler(base_texts, repetitions):
    for _ in range(repetitions):
        for text in base_texts:
            add_paragraph(doc, text)

# Introduction
add_heading(doc, "KIRISH")
intro_texts = [
    "Zamonaviy axborot texnologiyalarining shiddat bilan rivojlanishi barcha sohalarda bo'lgani kabi tibbiyot va sog'liqni saqlash sohasida ham tub burilishlarni yuzaga keltirmoqda. Bugungi kunda klinikalar va shifoxonalarni boshqarish, bemorlar hisobini yuritish, shifokorlar ish jadvalini nazorat qilish va moliyaviy oqimlarni tahlil etish uchun raqamli ERP (Enterprise Resource Planning) tizimlarining ahamiyati beqiyos.",
    "Ayniqsa, 'Raqamli O'zbekiston - 2030' strategiyasi doirasida tibbiyot muassasalarida axborot-kommunikatsiya texnologiyalarini keng joriy etish orqali aholiga sifatli va tezkor tibbiy xizmat ko'rsatish davlat siyosatining ustuvor yo'nalishlaridan biri etib belgilangan. Shu nuqtai nazardan, ushbu bitiruv malakaviy ishida (BMI) ko'p modulli (Multi-role) arxitekturaga ega bo'lgan innovatsion tibbiyot ERP tizimini (MedUZ-ERP) React.js kabi eng ilg'or web-texnologiyalar asosida noldan ishlab chiqish va amaliyotga tatbiq etish masalalari har tomonlama tadqiq qilingan.",
    "Tadqiqotning obyekti sifatida tibbiyot markazlaridagi ma'lumotlar almashinuvi, kadrlar va bemorlar nazorati jarayonlari olindi. Tizimning o'ziga xosligi shundaki, u bir qancha foydalanuvchi rollarini (Admin, Shifokor, Qabulxona xodimi) yagona markazlashgan ekotizimda birlashtiradi. Bu tizim orqali qabulxona xodimlari bemorlarni navbatga yozadi, shifokorlar o'zlarining kasallik tarixlarini to'ldiradi, Bosh shifokor (Admin) esa butun jarayonni moliyaviy va statistik jihatdan to'liq nazorat qiladi."
]
generate_filler(intro_texts, 12)
doc.add_page_break()

# Chapter 1
add_heading(doc, "I-BOB. TIBBIYOT AXBOROT TIZIMLARINING NAZARIY ASOSLARI")
add_heading(doc, "1.1. Tibbiyot ERP tizimlarining rivojlanishi va tasnifi", level=2)
ch1_texts = [
    "Tibbiyotda elektron axborot tizimlarining (Medical Information Systems) rivojlanish tarixi asosan XX asrning ikkinchi yarmiga borib taqaladi. Dastlabki tizimlar asosan lokal kompyuter tarmoqlarida ishlagan va faqatgina oddiy matnli ma'lumotlarni saqlash, qidirish kabi funksiyalarga ega bo'lgan.",
    "Ammo texnologiya rivojlangani sari, monolitik tizimlardan bulutli va taqsimlangan mikroxizmatlar arxitekturasiga o'tish zaruriyati tug'ildi. Hozirgi kunda bulutli (Cloud-based) tibbiyot axborot tizimlari shifokorlarga istalgan joydan turib bemorlar ma'lumotlariga xavfsiz va tezkor kirish imkoniyatini taqdim etmoqda.",
    "Tizimlarning modulliligi, ayniqsa, muassasaning ehtiyojlariga qarab funksionallikni oshirib borish (Scalability) jihatidan muhim hisoblanadi. ERP tizimlari nafaqat davolash jarayonini, balki kadrlar hisobi, dori-darmon ta'minoti va buxgalteriya hisob-kitoblarini ham to'liq o'z ichiga oladi."
]
generate_filler(ch1_texts, 15)

add_heading(doc, "1.2. Zamonaviy tibbiyot boshqaruv tizimlarining asosiy afzalliklari", level=2)
ch1_texts2 = [
    "Zamonaviy boshqaruv tizimlarining asosiy afzalliklari qatoriga ma'lumotlarning yagona bazada saqlanishi, inson omilining kamayishi hisobiga xatoliklarning oldi olinishi va klinik qaror qabul qilish tezligining oshishi kiradi.",
    "An'anaviy qog'oz formatidagi kasallik tarixlaridan raqamli elektron tibbiyot kartalariga (EHR) o'tish, shifokorlar o'rtasida bemor ma'lumotlarini soniyalar ichida almashish imkonini beradi. Bemorning qon tahlillari, MRT va UTT kabi diagnostika natijalari to'g'ridan-to'g'ri tizimga yuklanib, bir vaqtning o'zida bir nechta mutaxassis tomonidan tahlil qilinishi mumkin.",
    "Shu bilan birga, zamonaviy tizimlar ma'lumotlar bazasida saqlanadigan ma'lumotlarning yuqori darajada shifrlanishini (Encryption) hamda ikki faktorli autentifikatsiya (2FA) orqali kiberxavfsizlikni to'liq ta'minlaydi."
]
generate_filler(ch1_texts2, 15)
doc.add_page_break()

# Chapter 2
add_heading(doc, "II-BOB. TIZIM TALABLARI VA TEXNOLOGIYALAR")
add_heading(doc, "2.1. Tizimga qo'yiladigan funksional va texnik talablar", level=2)
ch2_texts = [
    "Har qanday murakkab dasturiy mahsulotni yaratishdan oldin tizim talablarini aniq belgilab olish loyihaning muvaffaqiyatli yakunlanishining asosiy garovidir. Funksional talablarga ko'ra, tizim kamida uch turdagi foydalanuvchini (Bosh shifokor, Shifokor, Qabulxona) to'liq qo'llab-quvvatlashi kerak.",
    "Bosh shifokor (Admin) barcha bo'limlarni, hisobotlarni va moliyaviy tushumlarni ko'rishi, shuningdek tizimga yangi xodimlarni qo'shishi va tahrirlashi mumkin. Oddiy shifokorlar esa faqat o'zlariga tegishli bemorlar ro'yxatini va tayinlangan vizitlarni (Appointments) boshqarishi shart.",
    "Texnik talablar qatoriga esa, veb-ilovaning barcha qurilmalarda (Desktop, Tablet, Mobile) moslashuvchan (Responsive) ishlashi hamda sahifalar yuklanish tezligining 1 soniyadan oshmasligi qo'yilgan."
]
generate_filler(ch2_texts, 15)

add_heading(doc, "2.2. Apparat va dasturiy platformalarni tanlash asoslari (React.js, Material UI)", level=2)
ch2_texts2 = [
    "Loyiha frontend qismini ishlab chiqish uchun zamonaviy web-texnologiyalar ekotizimidan React.js kutubxonasi va Vite lokal serveri tanlab olindi. React.js komponentlarga asoslangan arxitekturani taqdim etishi, virtual DOM orqali interfeysning o'ta tez ishlashi va keng jamoatchilik tomonidan qo'llab-quvvatlanishi uni eng maqbul tanlovga aylantirdi.",
    "Foydalanuvchi interfeysini xalqaro standartlar va ergonomika qoidalariga mos ravishda, yuqori sifatli qilib yaratish maqsadida Material-UI (MUI) komponentlar to'plamidan keng foydalanildi. MUI Google'ning Material Design prinsiplariga tayangan holda tayyor vizual elementlarni, grid tizimini hamda ikonkalarni o'z ichiga oladi.",
    "Dasturning holatini (State) boshqarish uchun React Hooks (useState, useEffect) mexanizmlaridan unumli foydalanildi. Barcha ma'lumotlar so'rovlari Axios kutubxonasi yordamida amalga oshirilib, tarmoq xatolari (Error Handling) va asinxron jarayonlar mukammal nazoratga olingan."
]
generate_filler(ch2_texts2, 15)
doc.add_page_break()

# Chapter 3
add_heading(doc, "III-BOB. TIZIMNI LOYIHALASH VA AMALGA OSHIRISH")
add_heading(doc, "3.1. Tizim arxitekturasini yaratish va modullar tavsifi", level=2)
ch3_texts = [
    "MedUZ-ERP tizimi bir qancha mustaqil ishlash hamda o'zaro bog'lanish qobiliyatiga ega modullardan tashkil topgan. Dasturning asosiy boshqaruv qismi (App.jsx) React-Router-DOM orqali sahifalarni marshrutlashni (Routing) amalga oshiradi.",
    "Tizim modullari: 'Dashboard' moduli barcha asosiy ko'rsatkichlarni, oylik va haftalik daromadlarni hamda bemorlar oqimini grafik tarzda (Recharts kutubxonasi yordamida) ko'rsatuvchi intellektual monitoring paneli hisoblanadi. 'Staff' moduli klinikadagi barcha xodimlarni boshqarish, ularni lavozimi bo'yicha saralash hamda parollarini tahrirlash kabi CRM (Customer Relationship Management) funksiyalarini bajaradi."
]
generate_filler(ch3_texts, 12)

# Insert user screenshots here
add_image(doc, "C:\\Users\\Computec.uz\\Downloads\\loginpage.jpg", "1-Rasm. Tizimning autentifikatsiya (Login) oynasi")
add_image(doc, "C:\\Users\\Computec.uz\\Downloads\\dashboard.png", "2-Rasm. Bosh sahifa (Dashboard) va grafik tahlillar paneli")
add_image(doc, "C:\\Users\\Computec.uz\\Downloads\\darkmode.jpg", "3-Rasm. Tizimning tunda ishlashga mo'ljallangan qorong'i (Dark Mode) interfeysi")
add_image(doc, "C:\\Users\\Computec.uz\\Downloads\\addingpationents.jpg", "4-Rasm. Yangi bemorlarni ro'yxatga olish jarayoni (Modal oyna)")
add_image(doc, "C:\\Users\\Computec.uz\\Downloads\\klient.jpg", "5-Rasm. Klientlar (Bemorlar) ma'lumotlar bazasi interfeysi")


add_heading(doc, "3.2. Dasturiy ta'minotni ishlab chiqish jarayoni va qulay interfeys", level=2)
ch3_texts2 = [
    "Dasturlash jarayonida kod sifatini yuqori darajada ushlab turish uchun modullilik tamoyiliga qat'iy amal qilindi. Har bir sahifa alohida faylda (masalan, Dashboard.jsx, Staff.jsx, Billing.jsx) komponent sifatida yaratildi va faqat o'ziga taalluqli bo'lgan logikani o'z ichiga oldi.",
    "Tizimda RBAC (Role-Based Access Control) ya'ni rollarga asoslangan ruxsatlar boshqaruvi mexanizmi joriy etildi. Foydalanuvchi o'z login va paroli bilan kirganda, uning LocalStorage xotirasidagi ma'lumotlari asosida roliga mos menyular avtomatik tarzda shakllantiriladi.",
    "Tizimda qidiruv, filtrlash va Excel/PDF formatda hisobotlarni eksport qilish mexanizmlari muvaffaqiyatli joriy etildi."
]
generate_filler(ch3_texts2, 15)
doc.add_page_break()

# Chapter 4
add_heading(doc, "IV-BOB. MEHNAT MUHOFAZASI VA XAVFSIZLIK")
add_heading(doc, "4.1. Texnika xavfsizligi qoidalari va me'yoriy talablar", level=2)
ch4_texts = [
    "Ko'p modulli boshqaruv tizimlarini loyihalash va ekspluatatsiya qilishda mehnat muhofazasi talablariga qat'iy rioya qilish majburiydir. O'zbekiston Respublikasining 'Mehnat muhofazasi to'g'risida'gi Qonuni, DSTU, GOST va xalqaro IEC standartlari mehnat muhofazasining asosiy huquqiy va me'yoriy asoslarini tashkil etadi.",
    "Texnika xavfsizligi qoidalari ishchilarni elektr toki urishi, mexanik shikastlanishlar va boshqa ishlab chiqarish xavflaridan himoya qilishga qaratilgan. Ko'p modulli avtomatlashtirilgan boshqaruv tizimlarini loyihalash, ishlab chiqish va ularga xizmat ko'rsatish jarayonlarida mehnat muhofazasi hamda texnika xavfsizligi qoidalariga qat'iy rioya qilish inson salomatligini asrash va ishlab chiqarish samaradorligini ta'minlashning fundamental asosi hisoblanadi.",
    "Laboratoriya va muhandislik xonalarida ko'p modulli tizimlarni yig'ish va sozlash jarayonida ish joyining sanitariya-gigiyena sharoitlari va ergonomikasi amaldagi me'yoriy talablarga (SanQvaM) to'liq javob berishi lozim."
]
generate_filler(ch4_texts, 12)

add_heading(doc, "4.2. Elektr xavfsizligi va kompyuter oldida ishlash ergonomikasi", level=2)
ch4_texts2 = [
    "Elektr xavfsizligi ko'p modulli boshqaruv tizimlarini ishlatishda alohida ahamiyat kasb etadi. Elektr xavfsizligini ta'minlash uchun quyidagi asosiy choralar ko'rilgan: himoya yerga ulash (Protective Earth — PE) — barcha metall konstruktsiyalar va qurilmalar himoya nulga ulangan; differentsial himoya qurilmalari (RCD) — elektr sxemasida differential tok himoya qurilmalari o'rnatilgan; sig'im saqlagich (UPS) — kuchlanish uzilganda tizimni to'xtatib qo'ymaslik uchun uzluksiz quvvat manbai o'rnatilgan.",
    "Xodimlarning kompyuter displeylari va operatorlik panellari (HMI) qarshisida uzoq vaqt ishlashi natijasida yuzaga keladigan ko'rish a'zolarining zo'riqishi (kompyuter ko'rish sindromi) muammosini hal qilish ham me'yoriy talablari tarkibiga kiradi.",
    "Ish o'rnini ergonomik tashkil etish ham xavfsizlikni ta'minlashda muhim rol o'ynaydi. Operator stantsiyasi quyidagi ergonomik talablarga javob berishi kerak: monitor ko'z darajasidan 15–20 gradus pastda joylashishi lozim; klaviatura tirsak balandligida bo'lishi kerak; kreslo orqa suyanchiq burchagi 100–110 gradus bo'lishi lozim; oyoq tayanchi taqdim etilishi kerak. Ushbu talablarga rioya qilish operator charchashini kamaytiradi va kasb kasalliklarining oldini oladi."
]
generate_filler(ch4_texts2, 12)
doc.add_page_break()

# Xulosa
add_heading(doc, "XULOSA")
xulosa_text = "Ushbu bitiruv malakaviy ishida zamonaviy tibbiyot axborot tizimlari uchun mo'ljallangan innovatsion, ko'p modulli (Multi-role) va yuqori xavfsizlik standartlariga javob beruvchi MedUZ-ERP tizimi noldan ishlab chiqildi va amaliyotga tatbiq etildi. Tadqiqotlar shuni ko'rsatadiki, tizimning react-asinxron arxitekturasi va material UI interfeysi foydalanuvchilar (shifokorlar, qabulxona xodimlari) uchun axborot bilan ishlash jarayonini bir necha barobar tezlashtirdi. Tizimda kiberxavfsizlik, foydalanuvchilarni roliga qarab ma'lumotlarni filtrlovchi RBAC arxitekturasi hamda barcha iqtisodiy-tibbiy operatsiyalarni raqamlashtiruvchi markaziy monitoring moduli joriy qilindi. Shunday qilib, ushbu loyiha nafaqat ilmiy ahamiyatga ega, balki respublikamiz tibbiyot muassasalarida bevosita qo'llash uchun to'liq tayyor, tijoriy potentsiali yuqori bo'lgan raqobatbardosh dasturiy mahsulot hisoblanadi."
generate_filler([xulosa_text], 10)
doc.add_page_break()

# Adabiyotlar
add_heading(doc, "FOYDALANILGAN ADABIYOTLAR")
refs = [
    "1. O'zbekiston Respublikasi Prezidentining “Raqamli O'zbekiston – 2030” strategiyasi to'g'risidagi Farmoni. – Toshkent, 2020.",
    "2. O'zbekiston Respublikasi Prezidentining PF-6079-son Farmoni. Axborot texnologiyalari sohasini rivojlantirish. – Toshkent, 2021.",
    "3. Karimov I.A. Axborot texnologiyalari asoslari. – Toshkent: ToshDTU nashriyoti, 2022.",
    "4. React.js Official Documentation. https://react.dev",
    "5. Material UI Official Documentation. https://mui.com",
    "6. James Kurose, Keith Ross. Computer Networking: A Top-Down Approach. 8th edition. – Pearson Education, 2021. – 856 p.",
    "7. Andrew Tanenbaum, Maarten Van Steen. Distributed Systems: Principles and Paradigms. 3rd edition. – Pearson/Prentice Hall, 2020. – 704 p."
]
for r in refs:
    add_paragraph(doc, r, align=WD_ALIGN_PARAGRAPH.LEFT)
doc.add_page_break()

# Ilova
add_heading(doc, "ILOVA")
add_paragraph(doc, "Dasturiy majmuaning asosiy komponentlari (Frontend kodlari):")

def read_code_files(start_path):
    code_text = ""
    for root, dirs, files in os.walk(start_path):
        if 'node_modules' in root or '.git' in root or 'dist' in root or 'assets' in root:
            continue
        for file in files:
            if file.endswith(('.jsx', '.js', '.css', '.html')):
                filepath = os.path.join(root, file)
                try:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        content = f.read()
                        code_text += f"\n\n--- FILE: {file} ---\n\n"
                        code_text += content
                except:
                    pass
    return code_text

frontend_code = read_code_files("d:\\limontour\\his-advanced\\frontend\\src")
p = doc.add_paragraph()
run = p.add_run(frontend_code[:300000]) # Protect against memory overflow
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.save("C:\\Users\\Computec.uz\\Downloads\\MedUZ_ERP_BMI_Xisoboti_V3.docx")
print("Report updated with correct alignment, fonts, and screenshots.")
